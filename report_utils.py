"""
Report Utilities Module

This module contains common helper functions used by both the MCP and Data Chiefs
report generation scripts. This helps to avoid code duplication and ensures
consistent functionality across both report types.
"""
import os
import sys
from datetime import datetime
import re
import ast
import json
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def resource_path(relative_path):
    """
    Get absolute path to resource, works for dev and PyInstaller.
    
    This function helps finding resources whether the script is run directly
    or from a bundled executable created with PyInstaller.
    """
    try:
        base_path = sys._MEIPASS
    except AttributeError:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

# --- Document Handling Functions ---
def _safe_get_table(doc, table_index, default=None):
    """
    Safely retrieves a table, returning default if not found.
    
    Args:
        doc: The document object
        table_index: Index of the table to retrieve
        default: Value to return if table doesn't exist
        
    Returns:
        The table object or default value
    """
    try: 
        return doc.tables[table_index]
    except IndexError:
        print(f"Warning: Table {table_index} not found.")
        return default

def _safe_get_cell(table, row_index, col_index, default=None):
    """
    Safely retrieves a cell, returning default if not found.
    
    Args:
        table: The table object
        row_index: Row index
        col_index: Column index
        default: Value to return if cell doesn't exist
        
    Returns:
        The cell object or default value
    """
    try: 
        return table.cell(row_index, col_index)
    except IndexError:
        print(f"Warning: Cell ({row_index}, {col_index}) not found.")
        return default

def _safe_set_text(cell, text):
    """
    Safely sets cell text, clearing existing content.
    
    Args:
        cell: The cell object
        text: Text content to set
    """
    if cell:
        for p in cell.paragraphs:
            p = p._element
            p.getparent().remove(p)
        paragraph = cell.add_paragraph()
        run = paragraph.add_run(str(text))
        run.font.name = 'Montserrat Light'
        run.font.size = Pt(10)

def _safe_add_paragraph(cell, text):
    """
    Safely adds a paragraph to a cell with proper formatting.
    
    Args:
        cell: The cell object
        text: Text content to add
    """
    if cell:
        paragraph = cell.add_paragraph(text)
        run = paragraph.runs[0]
        run.font.name = 'Montserrat'
        run.font.size = Pt(10)

        r = run._element
        rPr = r.rPr
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            r.append(rPr)

        rFonts = OxmlElement('w:rFonts')
        rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii', 'Montserrat Light')
        rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi', 'Montserrat Light')
        rPr.append(rFonts)

def _safe_literal_eval(s, default=None):
    """
    Safely evaluates a string as a Python literal, removing backslashes.
    
    Args:
        s: String to evaluate
        default: Value to return if evaluation fails
        
    Returns:
        Evaluated Python object or default value
        
    Note:
        This improved version replaces 'N/A' with None instead of -99
        for better clarity and safety.
    """
    try:
        s = s.replace("\\", "")
        # Replace 'N/A' with None instead of -99 for better clarity
        s = s.replace("'N/A'", "None")
        s = s.replace('"N/A"', "None")
        return ast.literal_eval(s)
    except (SyntaxError, ValueError) as e:
        print(f"Error: Error evaluating string: {s} - {e}")
        return default

def format_bullet_points(text):
    """
    Formats text with bullet points to ensure proper bullet characters (* -> •).
    It retains single newlines for separation, which add_bulleted_content will handle.
    
    Args:
        text: Text containing bullet points (either with * or •)
        
    Returns:
        Formatted text with • bullet points and single newline separators.
    """
    if not isinstance(text, str):
        return text
        
    # Split the text into lines
    lines = text.split('\n')
    formatted_lines = []
    
    for line in lines:
        line = line.strip()
        if not line:
            # Keep single empty lines if they exist between points, but don't add extra
            if formatted_lines and formatted_lines[-1]: # Check if the last added line wasn't empty
                 formatted_lines.append('')
            continue
            
        # Check if this is a bullet point (either * or •)
        if line.startswith('*') or line.startswith('•'):
            # Remove the bullet and any leading spaces
            content = line[1:].strip()
            # Add back the bullet with proper spacing
            formatted_lines.append(f'• {content}')
        else:
            # Regular paragraph, keep as is
            formatted_lines.append(line)
    
    # Join lines with single newlines
    return '\n'.join(formatted_lines)

def shuttle_text(shuttle):
    """Helper function to get combined text from a list of runs."""
    t = ''
    for run in shuttle:
        t += run.text
    return t

def replace_text_preserving_format(doc, data):
    """
    Replaces text in paragraphs and tables, preserving formatting.
    Handles cases where the text to replace spans multiple runs.
    Args:
        doc: The python-docx Document object.
        data: A dictionary {key_to_replace: replacement_value}.
              Replacement value may contain '<<BREAK>>' markers.
    """
    print("Replacing text while preserving format...")
    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    # Also include header/footer paragraphs if necessary
    for section in doc.sections:
        paragraphs.extend(section.header.paragraphs)
        paragraphs.extend(section.footer.paragraphs)
        # Add paragraphs from header/footer tables if needed
        for table in section.header.tables:
             for row in table.rows:
                 for cell in row.cells:
                     paragraphs.extend(cell.paragraphs)
        for table in section.footer.tables:
             for row in table.rows:
                 for cell in row.cells:
                     paragraphs.extend(cell.paragraphs)


    for key, value in data.items():
        key_to_find = str(key) # Placeholder like {prompt3_personality}
        replacement_value = str(value)

        for p in paragraphs:
            if key_to_find not in p.text:
                continue

            # This inner logic tries to find the key across runs
            begin = 0
            while begin < len(p.runs):
                end = begin
                current_text = ""
                key_found_in_shuttle = False

                # Expand the shuttle until the key is potentially found
                while end < len(p.runs):
                    shuttle = p.runs[begin:end+1]
                    current_text = shuttle_text(shuttle)
                    if key_to_find in current_text:
                        key_found_in_shuttle = True
                        break
                    # If key starts within this shuttle but isn't complete, keep expanding
                    partial_match = False
                    for i in range(len(key_to_find), 0, -1):
                         if current_text.endswith(key_to_find[:i]):
                              partial_match = True
                              break
                    if not partial_match and not key_to_find.startswith(current_text):
                         # Optimization: if key cannot start with current text, advance 'begin' faster
                          break
                    end += 1

                if key_found_in_shuttle:
                    # Key found spanning runs from 'begin' to 'end'
                    shuttle = p.runs[begin:end+1]
                    full_shuttle_text = shuttle_text(shuttle)
                    # print(f"Found '{key_to_find}' in runs {begin}-{end}: {[r.text for r in shuttle]}")

                    # Perform the replacement
                    start_index_in_full = full_shuttle_text.find(key_to_find)
                    end_index_in_full = start_index_in_full + len(key_to_find)

                    # Calculate which part belongs to which run and replace/clear
                    processed_len = 0
                    first_run_processed = False
                    for i, run in enumerate(shuttle):
                        run_len = len(run.text)
                        run_start = processed_len
                        run_end = processed_len + run_len

                        # Determine intersection of run with the key's location
                        replace_start_in_run = max(0, start_index_in_full - run_start)
                        replace_end_in_run = min(run_len, end_index_in_full - run_start)

                        if replace_start_in_run < replace_end_in_run: # This run overlaps with the key
                            original_text = run.text
                            if not first_run_processed:
                                # First run gets the replacement value + surrounding text
                                run.text = original_text[:replace_start_in_run] + replacement_value + original_text[replace_end_in_run:]
                                first_run_processed = True
                            else:
                                # Subsequent runs overlapping the key get cleared in that section
                                run.text = original_text[:replace_start_in_run] + original_text[replace_end_in_run:]
                        processed_len += run_len

                    # After replacement, restart search from the run *after* the replaced section
                    # This is tricky; a simpler approach might be to just advance 'begin' past 'end'
                    # or simply break and rely on multiple passes if needed.
                    # For simplicity, let's just advance begin past the affected runs.
                    begin = end + 1 # Move past the runs we just processed
                    continue # Continue the outer while loop

                else:
                     # Key not found starting at 'begin', advance 'begin'
                    begin += 1
    print("Text replacement finished.")

# Remove or comment out the old add_bulleted_content function
# def add_bulleted_content(doc, content, target_paragraph=None): ...

def split_paragraphs_and_apply_styles(doc):
    """
    Iterates through the document, splits paragraphs containing '<<BREAK>>',
    and applies 'List Bullet' style to lines starting with '•'.
    Must be called *after* all placeholders have been replaced.
    """
    print("Applying final paragraph splitting and styling...")
    # Need to iterate carefully as we modify the paragraph list
    paragraphs = list(doc.paragraphs) # Create a copy to iterate over
    for para in paragraphs:
        if '<<BREAK>>' in para.text:
            parts = para.text.split('<<BREAK>>')
            # Keep the first part in the current paragraph
            para.text = parts[0].strip()
            current_p_element = para._element # Reference point for inserting

            # Apply style to the first part if needed
            if para.text.startswith('•'):
                para.style = 'List Bullet'
                para.text = para.text[1:].strip() # Remove bullet character
            elif para.style.name.startswith('List Bullet'): # Ensure non-bullets don't keep bullet style
                para.style = 'Normal'

            # Insert new paragraphs for the remaining parts
            for part in parts[1:]:
                stripped_part = part.strip()
                if not stripped_part:
                    # Insert an empty paragraph for spacing
                    new_para = doc.add_paragraph('') # Add empty paragraph
                else:
                    # Insert paragraph with text
                    new_para = doc.add_paragraph(stripped_part)

                    # Apply bullet style if necessary
                    if stripped_part.startswith('•'):
                        new_para.text = stripped_part[1:].strip() # Remove bullet char
                        new_para.style = 'List Bullet'
                    else:
                         new_para.style = 'Normal' # Default style

                # Move the newly created paragraph right after the previous one
                current_p_element.addnext(new_para._element)
                current_p_element = new_para._element # Update reference point

            # Re-apply font to the original (now modified) paragraph and new ones
            all_paras_to_style = [para] + list(para._element.xpath('following-sibling::w:p'))[:len(parts)-1]
            for p_to_style in all_paras_to_style:
                 # Convert element back to Paragraph object if needed, or apply style via XML
                 # For simplicity, let's re-fetch the paragraph object if possible (might be slow)
                 # A better way involves direct XML font application or ensuring styles are correct.
                 try:
                    # Find the paragraph object corresponding to the element (this is non-trivial)
                    # Simplified: Apply font to runs assuming we have the Paragraph object
                    actual_para_obj = None
                    if p_to_style == para._element:
                        actual_para_obj = para
                    else:
                        # Find the object (inefficient way shown)
                        for p_search in doc.paragraphs:
                            if p_search._element == p_to_style:
                                actual_para_obj = p_search
                                break

                    if actual_para_obj:
                         for run in actual_para_obj.runs:
                             run.font.name = 'Montserrat'
                             run.font.size = Pt(10)
                 except Exception as e:
                     print(f"Warning: Could not apply font to split paragraph: {e}")

    print("Finished applying final styles.")

# --- Text Processing Functions ---
def clean(text):
    """
    Cleans input text by removing markdown and special characters.
    
    Args:
        text: Text to clean
        
    Returns:
        Cleaned text
    """
    return re.sub(r'[\【】`]|(```python)|(\*\*)', '', str(text)).strip() if isinstance(text, str) else text

def strip_extra_quotes(input_string):
    """
    Removes leading/trailing double quotes.
    
    Args:
        input_string: String to process
        
    Returns:
        String without leading/trailing quotes
    """
    if isinstance(input_string, str) and input_string.startswith('"') and input_string.endswith('"'):
        return input_string[1:-1]
    return input_string

def replacePiet(text, name, gender):
    """
    Replaces 'Piet' and handles gender-specific pronouns.
    
    Args:
        text: Text to process
        name: Name to replace 'Piet' with
        gender: 'M' or 'F' to determine pronoun replacement
        
    Returns:
        Processed text
    """
    if not isinstance(text, str):
        return ""

    text = text.replace("Piet", name.split()[0])
    text = re.sub(r'\bthe trainee\b', name.split()[0], text, flags=re.IGNORECASE)

    if gender == 'M':
        replacements = {
            "She": "He", "she": "he", "Her": "Him", "her": "him",
            "Hers": "His", "hers": "his", "Herself": "Himself", "herself": "himself"
        }
    elif gender == 'F':
        replacements = {
            "He": "She", "he": "she", "Him": "Her", "him": "her",
            "His": "Her", "his": "her", "Himself": "Herself", "himself": "herself"
        }
    else:
        return text

    for word, replacement in replacements.items():
        text = re.sub(r'\b' + re.escape(word) + r'\b', replacement, text)
    return text

def replace_piet_in_list(items_list, name, gender):
    """
    Replaces 'Piet' in each string item of a list.
    
    Args:
        items_list: List of items to process
        name: Name to replace 'Piet' with
        gender: 'M' or 'F' to determine pronoun replacement
        
    Returns:
        Processed list
    """
    if not isinstance(items_list, list):
        return items_list
    
    result = []
    for item in items_list:
        if isinstance(item, str):
            result.append(replacePiet(item, name, gender))
        else:
            result.append(item)
    return result

def restructure_date(date_str):
    """
    Restructures date string to DD-MM-YYYY format.
    
    Args:
        date_str: Date string to restructure
        
    Returns:
        Restructured date string or empty string if invalid
    """
    date_str = date_str.replace('/', '-')

    try:
        datetime.strptime(date_str, '%d-%m-%Y')
        return date_str
    except ValueError:
        try:
            date_obj = datetime.strptime(date_str, '%Y-%m-%d')
            return date_obj.strftime('%d-%m-%Y')
        except ValueError:
            return ''

def replace_and_format_header_text(doc, new_text):
    """
    Replaces header text and formats it with correct styling.
    
    Args:
        doc: The document object
        new_text: Text to replace placeholder with
        
    Returns:
        None
    """
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            if '***' in paragraph.text:
                paragraph.text = paragraph.text.replace('***', new_text)
                for run in paragraph.runs:
                    run.font.name = 'Montserrat'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(*(0xED, 0x6B, 0x55))
                    run.bold = True
                    run.italic = False
                    rFonts = OxmlElement('w:rFonts')
                    rFonts.set(qn('w:ascii'), 'Montserrat')
                    rFonts.set(qn('w:hAnsi'), 'Montserrat')
                    run._element.rPr.append(rFonts)

def clean_up(loc_dic):
    """
    Loads and cleans JSON data, handling Gemini variations.
    
    This is a unified version of the clean_up function used by both
    report generation scripts.
    
    Args:
        loc_dic: Path to JSON file
        
    Returns:
        Dictionary with cleaned data
    """
    try:
        with open(loc_dic, 'r', encoding='utf-8') as f:
            loaded_data = json.load(f)

        cleaned_data = {}
        for key, value in loaded_data.items():
            if isinstance(value, str):
                cleaned_value = clean(value.replace("\\", ""))
                # Normalize whitespace
                cleaned_value = ' '.join(cleaned_value.split())
                cleaned_data[key] = cleaned_value
            elif isinstance(value, list):
                cleaned_data[key] = [clean(item.replace("\\","")) if isinstance(item,str) else item for item in value]
            else:
                cleaned_data[key] = clean(value)

        return cleaned_data
    except (FileNotFoundError, json.JSONDecodeError) as e:
        print(f"Error: Error loading/cleaning JSON: {e}")
        return {}

def open_file(file_path):
    """
    Opens file based on OS.
    
    Args:
        file_path: Path to file to open
        
    Returns:
        None
    """
    if os.name == 'nt':  # Windows
        os.startfile(file_path)
    elif os.name == 'posix':  # macOS, Linux
        os.system(f'open "{file_path}"') 

def split_paragraphs_at_marker_and_style(doc):
    """
    Iterates through the document, splits paragraphs containing '<<BREAK>>',
    creates new paragraphs, and applies 'List Bullet' style to lines starting with '•'.
    Must be called *after* all placeholders have been replaced.
    """
    print("Applying final paragraph splitting and styling for <<BREAK>> markers...")
    # Iterate backwards through paragraphs to safely insert new ones
    # Using indices is safer when modifying the list of paragraphs
    i = len(doc.paragraphs) - 1
    while i >= 0:
        para = doc.paragraphs[i]
        if '<<BREAK>>' in para.text:
            parts = para.text.split('<<BREAK>>')
            # The last part stays in the current paragraph (or is the only part if <<BREAK>> is at end)
            para.text = parts[-1].strip()
            current_p_element = para._element  # Reference point for inserting

            # Always treat this segment as the conclusion: Normal style, no bullet
            para.style = 'Normal'
            # Strip any leading bullet character if present
            if para.text.startswith('•'):
                para.text = para.text[1:].strip()

            # Ensure consistent font and size for the last part
            for run in para.runs:
                run.font.name = 'Montserrat'
                run.font.size = Pt(10)

            # Insert new paragraphs for the preceding parts *before* the current one (in reverse order)
            for part in reversed(parts[:-1]):
                stripped_part = part.strip()
                # Create a new paragraph element (empty for now)
                new_p = OxmlElement('w:p')
                # Insert the new paragraph element *before* the current one
                current_p_element.addprevious(new_p)

                # Now create the Paragraph object to manipulate text and style
                # Need to find the paragraph associated with new_p (this is the hard part)
                # A workaround: add text directly via XML, apply style via XML
                if stripped_part:
                    # Apply style (List Bullet or Normal) via XML
                    pPr = OxmlElement('w:pPr')
                    pStyle = OxmlElement('w:pStyle')
                    style_name = 'List Bullet' if stripped_part.startswith('•') else 'Normal'
                    pStyle.set(qn('w:val'), style_name)
                    pPr.append(pStyle)
                    # If bullet style, add numbering properties (assuming ID 1 again)
                    if style_name == 'List Bullet':
                        numPr = OxmlElement('w:numPr')
                        ilvl = OxmlElement('w:ilvl'); ilvl.set(qn('w:val'), '0')
                        numId = OxmlElement('w:numId'); numId.set(qn('w:val'), '1') # Check this ID!
                        numPr.append(ilvl); numPr.append(numId)
                        pPr.append(numPr)
                    new_p.append(pPr)

                    # Add text run via XML
                    run_element = OxmlElement('w:r')
                    # Apply font within the run properties if needed
                    rPr_run = OxmlElement('w:rPr')
                    rFonts = OxmlElement('w:rFonts')
                    rFonts.set(qn('w:ascii'), 'Montserrat'); rFonts.set(qn('w:hAnsi'), 'Montserrat')
                    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '20') # Size in half-points (10pt = 20)
                    rPr_run.append(rFonts); rPr_run.append(sz)
                    run_element.append(rPr_run)
                    # Add the text element
                    t_element = OxmlElement('w:t')
                    t_element.text = stripped_part[1:].strip() if stripped_part.startswith('•') else stripped_part
                    run_element.append(t_element)
                    new_p.append(run_element)
                # else: inserting the empty <w:p> handles blank lines from consecutive <<BREAK>>

                # Update the reference element for the next insertion
                current_p_element = new_p


        i -= 1 # Move to the previous paragraph index

    print("Finished applying final styles for <<BREAK>> markers.") 