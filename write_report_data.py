import os
import sys
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import ast
import json
import re

# Import common functions from report_utils
from report_utils import (
    resource_path, _safe_get_table, _safe_get_cell, _safe_set_text, 
    _safe_add_paragraph, _safe_literal_eval,
    clean, strip_extra_quotes, clean_up, replacePiet, replace_piet_in_list,
    restructure_date, replace_and_format_header_text, open_file,
    replace_text_preserving_format, split_paragraphs_at_marker_and_style
)

# --- Constants ---
DETAILS_TABLE_INDEX = 0
COGCAP_TABLE_INDEX = 1
CONCLUSION_TABLE_INDEX = 2
HUMAN_SKILLS_START_TABLE = 4
HUMAN_SKILLS_TABLE_COUNT = 5
TECH_SKILLS_START_TABLE = 9
TECH_SKILLS_TABLE_COUNT = 5
DATA_TOOLS_TABLE_START = 15
DATA_TOOLS_ITEMS_PER_TABLE = 5
INTERESTS_TABLE_INDEX = 16
LANGUAGE_SKILLS_TABLE_INDEX = 14

def replace_placeholder_in_docx(doc, placeholder, replacement, font_name='Montserrat', font_size=10):
    """Replaces a placeholder in the document with custom font."""
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            inline = paragraph.runs
            for i in range(len(inline)):
                if placeholder in inline[i].text:
                    inline[i].text = inline[i].text.replace(placeholder, replacement)
                    inline[i].font.name = font_name
                    inline[i].font.size = Pt(font_size)

def update_document(output_dic, name, assessor, gender, program):
    """Updates the Word document."""
    try:
        doc = Document(resource_path('resources/Assessment_report_Data_chiefs.docx'))
    except Exception as e:
        print(f"Error: Failed to open template: {e}")
        return None

    # --- Prepare Replacement Dictionary ---
    replacements = {}

    # Static replacements
    replacements['***'] = name.split()[0]
    replacements['ASSESSOR'] = assessor.upper()

    # Dynamic Content replacements
    dynamic_prompts = [
        'prompt2_firstimpr', 'prompt3_personality',
        'prompt4_cogcap_remarks'
        # Interests (prompt9) are handled separately via add_interests_table
    ]
    for prompt_key in dynamic_prompts:
        replacement_text = output_dic.get(prompt_key, "")
        if prompt_key in ['prompt2_firstimpr', 'prompt3_personality', 'prompt4_cogcap_remarks']:
            replacement_text = replacePiet(replacement_text, name, gender)
        replacements[f"{{{prompt_key}}}"] = replacement_text

    # Language Skill placeholders (assuming they exist in the Data template too)
    language_replacements_str = output_dic.get('prompt5_language', "[]")
    language_levels = _safe_literal_eval(language_replacements_str, [])
    if isinstance(language_levels, list):
        language_names = ["Dutch", "French", "English"]
        for index, language_name in enumerate(language_names):
            if index < len(language_levels):
                proficiency_level = language_levels[index]
                placeholder = f"{{prompt5_language_{language_name.lower()}}}"
                replacements[placeholder] = proficiency_level
            else:
                print(f"Warning: No proficiency level provided for {language_name}.")
                placeholder = f"{{prompt5_language_{language_name.lower()}}}"
                replacements[placeholder] = "N/A"

    # --- Perform ALL Text Replacements ---
    replace_text_preserving_format(doc, replacements)

    # --- Handle list prompts that may contain "Piet" ---
    # Operate on the _original JSON data for these prompts
    list_prompt_keys_original = ['prompt6a_conqual_original', 'prompt6b_conimprov_original']
    for original_key in list_prompt_keys_original:
        if original_key in output_dic:
            list_str = output_dic.get(original_key, "[]")
            # Safely evaluate the ORIGINAL JSON string
            list_items = _safe_literal_eval(list_str, [])
            if isinstance(list_items, list):
                # Replace Piet in each list item
                list_items_pietless = replace_piet_in_list(list_items, name, gender)
                # Store the processed list back into the _original key
                output_dic[original_key] = list_items_pietless # Store the list directly
            else:
                print(f"Warning: Could not process {original_key} as a list after eval.")
                output_dic[original_key] = [] # Ensure it's an empty list on failure
        else:
             # Ensure the key exists even if the prompt failed, to avoid errors later
             output_dic[original_key] = []

    # --- Table/Specific Location Content ---
    add_content_detailstable(doc, [name, "", program, "", ""])
    replace_and_format_header_text(doc, name)
    add_content_cogcaptable(doc, output_dic.get('prompt4_cogcap_scores', "[]"))
    
    # --- Add language levels to language skills table (14th table) ---
    language_replacements_str = output_dic.get('prompt5_language', "[]")
    # Ensure backslashes are removed before parsing
    if isinstance(language_replacements_str, str):
        language_replacements_str = language_replacements_str.replace("\\", "")
    language_levels = _safe_literal_eval(language_replacements_str, [])
    update_language_skills_table(doc, language_levels)

    # --- Conclusion Table ---
    # Pass the processed list from the _original key
    conclusion(doc, 0, output_dic.get('prompt6a_conqual_original', []))
    conclusion(doc, 1, output_dic.get('prompt6b_conimprov_original', []))

    # --- Interests ---
    interests_str = output_dic.get('prompt9_interests', "")
    add_interests_table(doc, interests_str)

    # Profile review (icons)
    qual_scores_str = output_dic.get('prompt7_qualscore_data', "[]")
    qual_scores = _safe_literal_eval(qual_scores_str, [])
    if isinstance(qual_scores, list) and len(qual_scores) >= 23:
        add_icons_data_chief(doc, qual_scores[:18])
        add_icons_data_chief_2(doc, qual_scores[18:23])
    else:
        print(f"Warning: Invalid qual_scores data.")

    # Data tools (icons)
    data_tools_str = output_dic.get('prompt8_datatools', "[]")
    data_tools_scores = _safe_literal_eval(data_tools_str, [])
    if isinstance(data_tools_scores, list):
        add_icons_data_tools(doc, data_tools_scores)
    else:
        print(f"Warning: Invalid data_tools_scores data.")

    # --- Save Document ---
    current_time = datetime.now()
    formatted_time = current_time.strftime("%m%d%H%M")
    
    # Define output directory and ensure it exists
    output_dir = "output_reports"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    # Save to the output directory
    updated_doc_path = os.path.join(output_dir, f"Assessment Report - {name} - {formatted_time}.docx")
    try:
        # Apply final paragraph splitting and styling *before* saving
        split_paragraphs_at_marker_and_style(doc) # This handles the display format
        doc.save(updated_doc_path)
        print(f"Document saved: {updated_doc_path}") # Added print statement
        return updated_doc_path
    except Exception as e:
        print(f"Error: Failed to save document: {e}")
        return None

def format_datatools_output(datatools_json_string):
    """Formats data tools output from JSON string."""
    try:
        datatools_dict = ast.literal_eval(datatools_json_string)
        formatted_text = ""
        for tool, level in datatools_dict.items():
            formatted_text += f"- {tool}: {level}\n"
        return formatted_text.strip()
    except (ValueError, SyntaxError):
        return "Could not parse data tools information."


def format_interests_output(interests_json_string):
    """Formats interests output from JSON string."""
    try:
        # Clean the string by removing backslashes
        interests_json_string = interests_json_string.replace("\\", "")
        # Replace 'N/A' with a placeholder if it's the only item
        if interests_json_string.strip() == '"N/A"' or interests_json_string.strip() == "'N/A'":
            return "No specific interests identified"
            
        interests_list = ast.literal_eval(interests_json_string)
        
        # If the list contains only 'N/A', return a placeholder
        if interests_list == ['N/A'] or interests_list == ["N/A"]:
            return "No specific interests identified"
            
        formatted_text = ""
        for interest in interests_list:
            # Skip 'N/A' entries
            if interest == 'N/A' or interest == "N/A":
                continue
            formatted_text += f"- {interest}\n"
            
        # If no valid interests were found, return a placeholder
        if not formatted_text:
            return "No specific interests identified"
            
        return formatted_text.strip()
    except (ValueError, SyntaxError) as e:
        print(f"Error parsing interests: {e}")
        return "Could not parse interests information."

def add_content_detailstable(doc, personal_details):
    """Adds personal details to the first table."""
    table = _safe_get_table(doc, DETAILS_TABLE_INDEX)
    if not table:
        return

    if not isinstance(personal_details, list):
        print(f"Warning: personal_details is not a list.")
        return

    if len(personal_details) == 1 and all(isinstance(ele, str) for ele in personal_details):
        personal_details = personal_details[0].split(',')

    for row_index, row in enumerate(table.rows):
        if len(row.cells) > 1:
            first_cell_text = row.cells[0].text.strip()
            second_cell_text = row.cells[1].text.strip()

            if first_cell_text == "Name candidate" and second_cell_text == ":":
                cell = _safe_get_cell(table, row_index, 2)
                _safe_set_text(cell, personal_details[0] if len(personal_details) > 0 else '')

            if first_cell_text == "Date of birth" and second_cell_text == ":":
                cell = _safe_get_cell(table, row_index, 2)
                _safe_set_text(cell, restructure_date(personal_details[1]) if len(personal_details) > 1 else '')

            if first_cell_text == "Position" and second_cell_text == ":":
                cell = _safe_get_cell(table, row_index, 2)
                _safe_set_text(cell, personal_details[2] if len(personal_details) > 2 else '')

            if first_cell_text == "Assessment date" and second_cell_text == ":":
                cell = _safe_get_cell(table, row_index, 2)
                _safe_set_text(cell, restructure_date(personal_details[3]) if len(personal_details) > 3 else '')

            if first_cell_text == "Pool" and second_cell_text == ":":
                cell = _safe_get_cell(table, row_index, 2)
                _safe_set_text(cell, personal_details[4] if len(personal_details) > 4 else '')

def add_content_cogcaptable(doc, scores_str):
    """Adds cognitive capacity scores."""
    table = _safe_get_table(doc, COGCAP_TABLE_INDEX)
    if not table:
        return

    scores = _safe_literal_eval(scores_str, [])
    if not isinstance(scores, list) or len(scores) != 6:
        print(f"Warning: Invalid scores data. Expected a list of 6 numbers.")
        return

    for i in range(6):
        cell = _safe_get_cell(table, 1, i + 1)
        if cell:
            if i == 0:
                _safe_set_text(cell, scores[i])
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0]
                run.bold = True
                run.underline = True
                paragraph.alignment = 1
            else:
                _safe_set_text(cell, scores[i])
                paragraph = cell.paragraphs[0]
                paragraph.alignment = 1



def add_content_cogcaptable_remark(doc, cogcap_output):
    """Adds remarks to the cognitive capacity table."""
    if not isinstance(cogcap_output, str):
        print(f"Warning: cogcap_output is not a string.")
        return

    table = _safe_get_table(doc, COGCAP_TABLE_INDEX)
    if not table:
        return

    remark_cell = _safe_get_cell(table, 2, 1)
    if not remark_cell:
        return

    _safe_set_text(remark_cell, cogcap_output)


def add_icons_data_chief(doc, list_scores):
    """Adds icons to Human Skills tables."""
    if not isinstance(list_scores, list):
        print(f"Warning: list_scores is not a list.")
        return

    score_index = 0
    for table_no_offset in range(HUMAN_SKILLS_TABLE_COUNT):
        table_no = HUMAN_SKILLS_START_TABLE + table_no_offset
        table = _safe_get_table(doc, table_no)
        if not table:
            continue

        for row_no in range(1, len(table.rows)):
            if score_index < len(list_scores):
                cell = _safe_get_cell(table, row_no, 0)
                if cell and cell.text.strip().startswith("AA"):
                    add_icon_to_cell(cell, list_scores[score_index])
                    score_index += 1
            else:
                # If we run out of scores, add N/A for remaining cells
                cell = _safe_get_cell(table, row_no, 0)
                if cell and cell.text.strip().startswith("AA"):
                    run = cell.paragraphs[0].add_run("N/A")
                    run.font.name = 'Montserrat'
                    run.font.size = Pt(9)

def add_icons_data_chief_2(doc, list_scores):
    """Adds icons to Technical Skills tables."""
    if not isinstance(list_scores, list):
        print(f"Warning: list_scores is not a list.")
        return

    score_index = 0
    for table_no_offset in range(TECH_SKILLS_TABLE_COUNT):
        table_no = TECH_SKILLS_START_TABLE + table_no_offset
        table = _safe_get_table(doc, table_no)
        if not table:
            continue

        for row_no in range(1, len(table.rows)):
            if score_index < len(list_scores):
                cell = _safe_get_cell(table, row_no, 0)
                if cell and cell.text.strip().startswith("AA"):
                    add_icon_to_cell(cell, list_scores[score_index])
                    score_index += 1
            else:
                # If we run out of scores, add N/A for remaining cells
                cell = _safe_get_cell(table, row_no, 0)
                if cell and cell.text.strip().startswith("AA"):
                    run = cell.paragraphs[0].add_run("N/A")
                    run.font.name = 'Montserrat'
                    run.font.size = Pt(9)

def add_icons_data_tools(doc, list_scores):
    """Adds icons to Data Tools tables."""
    if not isinstance(list_scores, list):
        print(f"Warning: list_scores is not a list.")
        return
    
    # Ensure we have exactly 5 scores, padding with None (our N/A placeholder) if needed
    if len(list_scores) < 5:
        list_scores = list_scores + [None] * (5 - len(list_scores))
    elif len(list_scores) > 5:
        list_scores = list_scores[:5]
    
    # Process each score
    for i in range(len(list_scores)):
        table_no = DATA_TOOLS_TABLE_START + (i // DATA_TOOLS_ITEMS_PER_TABLE)
        row_no = (i % DATA_TOOLS_ITEMS_PER_TABLE) + 2

        table = _safe_get_table(doc, table_no)
        if not table:
            continue

        cell = _safe_get_cell(table, row_no, 0)
        if cell:
            add_icon_to_cell(cell, list_scores[i])


def add_icon_to_cell(cell, score):
    """
    Adds an icon based on the score to a cell.
    
    This function has been updated to handle None values properly, which are
    now used to represent "N/A" instead of -99.
    """
    if cell is None:
        print(f"Warning: add_icon_to_cell called with None cell.")
        return
    
    _safe_set_text(cell, "")
    
    # Handle None (N/A) or non-integer scores
    if score is None or not isinstance(score, int):
        try:
            # Try to convert to int if possible (but not if None)
            score = int(score) if score is not None else None
        except (ValueError, TypeError):
            # If it's N/A or cannot be converted, use a default icon or text
            print(f"Warning: Non-integer score encountered: {score}. Using N/A.")
            run = cell.paragraphs[0].add_run("N/A")
            run.font.name = 'Montserrat'
            run.font.size = Pt(9)
            return

    # Special handling for None (our placeholder for N/A)
    if score is None:
        run = cell.paragraphs[0].add_run("N/A")
        run.font.name = 'Montserrat'
        run.font.size = Pt(9)
        return
        
    run = cell.paragraphs[0].add_run()
    if score == -1:
        run.add_picture(resource_path("resources/improvement.png"), width=Inches(.3))
    elif score == 0:
        run.add_picture(resource_path("resources/average.png"), width=Inches(.3))
    elif score == 1:
        run.add_picture(resource_path("resources/strong.png"), width=Inches(.3))
    else:
        print(f"Warning: Invalid score value: {score}")

def add_interests_table(doc, interests_text):
    """Fills in interests into the Interests Table as comma-separated text."""
    table = _safe_get_table(doc, INTERESTS_TABLE_INDEX)
    if not table:
        return

    # Handle the case where interests_text is 'N/A'
    if interests_text == 'N/A' or interests_text == "N/A":
        interests_string = "No specific interests identified"
    elif isinstance(interests_text, str):
        # Clean the string
        interests_text = interests_text.replace("\\", "")
        
        # If the entire string is just N/A in quotes
        if interests_text.strip() == '"N/A"' or interests_text.strip() == "'N/A'":
            interests_string = "No specific interests identified"
        else:
            # Process as a list
            try:
                interests_list = _safe_literal_eval(interests_text, [])
                
                # Filter out N/A values
                interests_list = [s for s in interests_list if s != 'N/A' and s != "N/A" and s is not None]
                
                if not interests_list:
                    interests_string = "No specific interests identified"
                else:
                    interests_string = ', '.join(interests_list)
            except Exception as e:
                print(f"Error processing interests: {e}")
                # Fallback to simple string processing if literal_eval fails
                interests_list = [s.strip() for s in interests_text.strip("[]").split(",") if s.strip()]
                interests_list = [s.strip('"').strip("'") for s in interests_list if s.strip('"').strip("'") != 'N/A']
                
                if not interests_list:
                    interests_string = "No specific interests identified"
                else:
                    interests_string = ', '.join(interests_list)
    else:
        print(f"Warning: interests_text is not a string.")
        interests_string = "No specific interests identified"

    cell = _safe_get_cell(table, 1, 0)
    if cell:
        _safe_set_text(cell, interests_string)
    else:
        print(f"Warning: Could not find cell to add interests text.")


def conclusion(doc, column, list_items):
    """Adds conclusion points (already processed list) to the specified column."""
    table = _safe_get_table(doc, CONCLUSION_TABLE_INDEX)
    if not table:
        return

    # Expecting list_items to be a Python list already
    if not isinstance(list_items, list):
        print(f"Warning: conclusion expected a list, got {type(list_items)}")
        return

    cell = _safe_get_cell(table, 1, column)
    if not cell:
        return
    # Clear cell content first
    _safe_set_text(cell, "")

    # Add each item as a separate paragraph with bullet formatting
    # We use _safe_add_paragraph which applies basic font, style is handled by cell/table
    for point in list_items:
        if isinstance(point, str):
             # Add pseudo-bullet for visual consistency within the table cell
            _safe_add_paragraph(cell, f'•  {point}')
        elif point: # Handle non-string items if necessary
            _safe_add_paragraph(cell, f'•  {str(point)}')


# Last style improvements
def replace_and_format_header_text(doc, new_text):
    """Replaces header text and formats it."""
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            if '***' in paragraph.text:
                paragraph.text = paragraph.text.replace('***', new_text)
                for run in paragraph.runs:
                    run.font.name = 'Montserrat'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(*(0xED, 0x6B, 0x55))
                    run.bold = True
                    run.italic = False
                    rFonts = OxmlElement('w:rFonts')
                    rFonts.set(qn('w:ascii'), 'Montserrat SemiBold')
                    rFonts.set(qn('w:hAnsi'), 'Montserrat SemiBold')
                    run._element.rPr.append(rFonts)

def update_language_skills_table(doc, language_levels):
    """
    Updates the language skills table (14th table) with language proficiency levels.
    
    Args:
        doc: The Word document
        language_levels: List of language levels [Dutch, French, English]
    """
    # Get the language skills table (14th table)
    table = _safe_get_table(doc, LANGUAGE_SKILLS_TABLE_INDEX)
    if not table:
        print("Warning: Language skills table not found.")
        return
    
    # Define language names for row identification
    language_names = ["Dutch", "French", "English"]
    
    # Valid language levels for matching
    valid_levels = ['A1', 'A2', 'B1', 'B2', 'C1', 'C2']
    
    # Find all rows that contain "A1/B1/B2.." - these are our language rows
    language_rows = []
    for row_index, row in enumerate(table.rows):
        # Skip header row
        if row_index == 0:
            continue
            
        # Get the first cell text to identify if it's a language row
        if len(row.cells) == 0:
            continue
            
        first_cell_text = row.cells[0].text.strip()
        if "A1/B1/B2" in first_cell_text:
            language_rows.append(row_index)
    
    # Update each language row with its corresponding level
    for i, row_index in enumerate(language_rows):
        if i >= len(language_levels):
            print(f"Warning: No level provided for language row {i+1}")
            continue
            
        raw_level = language_levels[i]
        
        # Normalize the level - extract A1/B1/C1 pattern if present
        normalized_level = None
        if isinstance(raw_level, str):
            # Clean any remaining quotes or backslashes
            raw_level = raw_level.replace('"', '').replace("'", "").replace("\\", "").strip()
            
            # Try to find a valid level pattern
            for valid_level in valid_levels:
                if valid_level.upper() in raw_level.upper():
                    normalized_level = valid_level.upper()
                    break
                    
            # If we couldn't find a match, look for level characters (A/B/C) and numbers (1/2)
            if normalized_level is None:
                level_match = re.search(r'([A-Ca-c]).*?([1-2])', raw_level)
                if level_match:
                    level_char = level_match.group(1).upper()
                    level_num = level_match.group(2)
                    normalized_level = f"{level_char}{level_num}"
                    
                    # Verify it's a valid level
                    if normalized_level not in [level.upper() for level in valid_levels]:
                        print(f"Warning: Extracted invalid level {normalized_level} from {raw_level}, using as is")
                    
        # If we couldn't normalize, use the raw level
        if normalized_level is None:
            print(f"Warning: Unable to normalize language level '{raw_level}', using as is")
            normalized_level = str(raw_level).upper()
        
        # Replace the "A1/B1/B2.." placeholder in the first cell
        row = table.rows[row_index]
        first_cell = row.cells[0]
        
        # Check if the first cell actually contains the A1/B1/B2 placeholder
        if "A1/B1/B2" in first_cell.text:
            # Extract any text before the placeholder (likely the language name)
            original_text = first_cell.text.strip()
            language_prefix = original_text.split("A1/B1/B2")[0].strip()
            
            # Set the text to include both the language name and the level
            _safe_set_text(first_cell, "")
            para = first_cell.paragraphs[0]
            
            # Add language name with original formatting
            if language_prefix:
                run_prefix = para.add_run(language_prefix + " ")
                run_prefix.font.name = 'Montserrat'
                run_prefix.font.size = Pt(10)
            
            # Add the level with bold formatting
            run_level = para.add_run(normalized_level)
            run_level.font.name = 'Montserrat'
            run_level.font.size = Pt(10)
            run_level.font.bold = True
        else:
            # Just set the level if we don't find the expected placeholder
            _safe_set_text(first_cell, normalized_level)
            run = first_cell.paragraphs[0].runs[0] if first_cell.paragraphs[0].runs else first_cell.paragraphs[0].add_run()
            run.font.name = 'Montserrat'
            run.font.size = Pt(10)
            run.font.bold = True
        
        # Find all cells with proficiency level placeholders (A1, B1, B2, etc.)
        for cell_index, cell in enumerate(row.cells):
            if cell_index == 0:  # Skip the first cell we just updated
                continue
                
            cell_text = cell.text.strip()
            
            # Check if this cell has an A1/B1/C1-style placeholder
            if re.match(r'[A-C][1-2]', cell_text):
                # Format the cell based on whether it matches the candidate's level
                if cell_text.upper() == normalized_level:
                    # Highlight the matched level
                    _safe_set_text(cell, "")
                    para = cell.paragraphs[0]
                    run = para.add_run(normalized_level)
                    run.font.name = 'Montserrat'
                    run.font.size = Pt(10)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Black
                else:
                    # Keep the placeholder for other levels, but make it less prominent
                    _safe_set_text(cell, "")
                    para = cell.paragraphs[0]
                    run = para.add_run(cell_text)
                    run.font.name = 'Montserrat'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(150, 150, 150)  # Light gray